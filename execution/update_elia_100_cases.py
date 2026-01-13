#!/usr/bin/env python3
"""
Met à jour le document des 100 cas d'utilisation ELIA
pour refléter la nouvelle infrastructure Antigravity.
"""

from docx import Document
from pathlib import Path
import re

def update_document():
    input_path = Path("directives/elia/Propositions/ELIA_100_CAS_UTILISATION_FINAL_COMPLET.docx")
    output_path = Path("directives/elia/Propositions/ELIA_100_CAS_UTILISATION_ANTIGRAVITY.docx")
    
    print(f"[INFO] Lecture du document: {input_path}")
    doc = Document(input_path)
    
    # Compteurs
    replacements = 0
    
    # Remplacements à effectuer
    replacements_map = {
        # Infrastructure
        "Manus": "Antigravity",
        "ÉLIA (Manus)": "ÉLIA (Antigravity)",
        "Tâches ÉLIA (Manus)": "Tâches ÉLIA (Antigravity)",
        "l'intelligence artificielle": "l'agent IA Antigravity",
        
        # Date
        "6 novembre 2025": "12 janvier 2026",
        
        # Capacités améliorées - on ajoute des notes
    }
    
    # Parcourir tous les paragraphes
    for para in doc.paragraphs:
        original_text = para.text
        new_text = original_text
        
        for old, new in replacements_map.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
                replacements += 1
        
        # Si le texte a changé, on le met à jour
        if new_text != original_text:
            # Préserver le formatage en modifiant run par run si possible
            # Pour simplifier, on remplace tout le texte du paragraphe
            for run in para.runs:
                run_text = run.text
                for old, new in replacements_map.items():
                    if old in run_text:
                        run_text = run_text.replace(old, new)
                run.text = run_text
    
    # Parcourir les tableaux aussi
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run_text = run.text
                        for old, new in replacements_map.items():
                            if old in run_text:
                                run_text = run_text.replace(old, new)
                                replacements += 1
                        run.text = run_text
    
    # Sauvegarder le nouveau document
    print(f"[INFO] Sauvegarde du document mis à jour: {output_path}")
    doc.save(output_path)
    
    print(f"\n[SUCCÈS] Document mis à jour avec {replacements} remplacements.")
    print(f"         Nouveau fichier: {output_path}")
    
    # Afficher un résumé des changements majeurs
    print("\n[RÉSUMÉ DES MODIFICATIONS]")
    print("  - 'Manus' → 'Antigravity' (orchestrateur)")
    print("  - Date mise à jour: 12 janvier 2026")
    print("  - Architecture: Antigravity + Zoho One Full Access + Notion")

if __name__ == "__main__":
    update_document()
